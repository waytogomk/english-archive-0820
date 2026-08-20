from __future__ import annotations

import base64
import csv
import hmac
import io
import json
import mimetypes
import os
import re
import sqlite3
import sys
import urllib.error
import urllib.request
import uuid
from datetime import datetime, timedelta, timezone
from http import HTTPStatus
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import parse_qs, urlparse

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parent
DATA_DIR = ROOT / "data"
UPLOAD_DIR = DATA_DIR / "uploads"
DB_PATH = DATA_DIR / "archive.db"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
ADMIN_SESSIONS: dict[str, datetime] = {}
ADMIN_SESSION_HOURS = 12


def now() -> str:
    return datetime.now(timezone.utc).isoformat()


def connect() -> sqlite3.Connection:
    connection = sqlite3.connect(DB_PATH)
    connection.row_factory = sqlite3.Row
    connection.execute("PRAGMA foreign_keys = ON")
    return connection


def init_db() -> None:
    with connect() as db:
        db.executescript(
            """
            CREATE TABLE IF NOT EXISTS contents (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              kind TEXT NOT NULL,
              title TEXT NOT NULL,
              school TEXT NOT NULL DEFAULT '',
              topic TEXT NOT NULL DEFAULT '',
              book TEXT NOT NULL DEFAULT '',
              source_role TEXT NOT NULL DEFAULT '',
              filename TEXT NOT NULL DEFAULT '',
              mime_type TEXT NOT NULL DEFAULT '',
              file_path TEXT NOT NULL DEFAULT '',
              extracted_text TEXT NOT NULL DEFAULT '',
              status TEXT NOT NULL DEFAULT 'draft',
              question_count INTEGER NOT NULL DEFAULT 0,
              version INTEGER NOT NULL DEFAULT 1,
              parent_id INTEGER,
              created_at TEXT NOT NULL,
              updated_at TEXT NOT NULL,
              FOREIGN KEY(parent_id) REFERENCES contents(id)
            );
            CREATE TABLE IF NOT EXISTS questions (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              content_id INTEGER NOT NULL,
              position INTEGER NOT NULL,
              prompt TEXT NOT NULL,
              choices_json TEXT NOT NULL DEFAULT '[]',
              answer TEXT NOT NULL,
              explanation TEXT NOT NULL DEFAULT '',
              FOREIGN KEY(content_id) REFERENCES contents(id) ON DELETE CASCADE
            );
            CREATE INDEX IF NOT EXISTS idx_contents_filter
              ON contents(kind, school, book, status, updated_at);
            """
        )


def safe_name(filename: str) -> str:
    clean = re.sub(r"[^A-Za-z0-9._-]+", "_", Path(filename).name)
    return clean[:120] or "upload.bin"


def extract_text(path: Path, mime_type: str) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf" or mime_type == "application/pdf":
        reader = PdfReader(path)
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        document = Document(path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            paragraphs.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
        return "\n".join(paragraphs)
    if suffix in {".txt", ".csv", ".tsv", ".json", ".md"} or mime_type.startswith("text/"):
        return path.read_text(encoding="utf-8-sig", errors="replace")
    if suffix in {".xlsx", ".xlsm"}:
        try:
            from openpyxl import load_workbook

            workbook = load_workbook(path, read_only=True, data_only=True)
            rows = []
            for sheet in workbook.worksheets:
                rows.append(f"[{sheet.title}]")
                rows.extend("\t".join("" if value is None else str(value) for value in row) for row in sheet.iter_rows(values_only=True))
            return "\n".join(rows)
        except Exception as exc:
            raise ValueError(f"Excel 파일을 읽지 못했습니다: {exc}") from exc
    return ""


def normalize_question(item: dict, position: int) -> dict | None:
    prompt = str(item.get("prompt") or item.get("question") or item.get("문제") or "").strip()
    answer = str(item.get("answer") or item.get("정답") or "").strip()
    if not prompt or not answer:
        return None
    choices = item.get("choices") or item.get("options") or item.get("보기") or []
    if isinstance(choices, str):
        choices = [part.strip() for part in re.split(r"\s*[|;/]\s*", choices) if part.strip()]
    explanation = str(item.get("explanation") or item.get("해설") or "").strip()
    return {"position": position, "prompt": prompt, "choices": choices, "answer": answer, "explanation": explanation}


def parse_structured_questions(text: str, suffix: str) -> list[dict]:
    questions: list[dict] = []
    if suffix == ".json":
        payload = json.loads(text)
        items = payload.get("questions", []) if isinstance(payload, dict) else payload
        for index, item in enumerate(items, start=1):
            if isinstance(item, dict) and (question := normalize_question(item, index)):
                questions.append(question)
        return questions
    if suffix in {".csv", ".tsv"}:
        dialect = "excel-tab" if suffix == ".tsv" else "excel"
        for index, item in enumerate(csv.DictReader(io.StringIO(text), dialect=dialect), start=1):
            if question := normalize_question(item, index):
                questions.append(question)
        return questions
    reading_questions = parse_reading_quiz_sections(text)
    return reading_questions or parse_questions_from_text(text)


def parse_reading_quiz_sections(text: str) -> list[dict]:
    """Import authored Reading Quiz sections without rewriting their content."""
    headers = list(re.finditer(r"(?im)^\s*Reading Quiz\s*\([^\n]+\)\s*$", text))
    if not headers:
        return []
    imported: list[dict] = []
    for header_index, header in enumerate(headers):
        end = headers[header_index + 1].start() if header_index + 1 < len(headers) else len(text)
        section = text[header.end():end].strip()
        parts = re.split(r"(?im)^\s*Answer Key\s*$", section, maxsplit=1)
        if len(parts) == 2:
            question_text, answer_text = parts
        else:
            answer_start = re.search(r"(?m)^\s*[A-D]\s+[—-]\s+", section)
            if not answer_start:
                continue
            question_text, answer_text = section[:answer_start.start()], section[answer_start.start():]
        question_matches = list(re.finditer(r"(?m)^\s*(\d{1,3})\.\s+", question_text))
        section_questions = []
        for question_index, match in enumerate(question_matches):
            block_end = question_matches[question_index + 1].start() if question_index + 1 < len(question_matches) else len(question_text)
            block = question_text[match.end():block_end].strip()
            option_matches = list(re.finditer(r"(?m)(?:^|\s)([A-D])\)\s*", block))
            if not option_matches:
                continue
            prompt = block[:option_matches[0].start()].strip()
            choices = []
            for option_index, option in enumerate(option_matches):
                option_end = option_matches[option_index + 1].start() if option_index + 1 < len(option_matches) else len(block)
                option_text = block[option.end():option_end].strip()
                choices.append(f"{option.group(1)}) {option_text}")
            section_questions.append({"prompt": prompt, "choices": choices})
        answer_lines = [line.strip() for line in answer_text.splitlines() if line.strip()]
        answers = []
        for line in answer_lines:
            if match := re.match(r"^([A-D])(?:\s*\((.*)\)|\s+[—-]\s+(.*))?\s*$", line, re.S):
                answers.append((match.group(1), (match.group(2) or match.group(3) or "").strip()))
        for question, answer in zip(section_questions, answers):
            imported.append(
                {
                    "position": len(imported) + 1,
                    "prompt": question["prompt"],
                    "choices": question["choices"],
                    "answer": answer[0],
                    "explanation": answer[1],
                }
            )
    return imported


def parse_questions_from_text(text: str) -> list[dict]:
    question_pattern = re.compile(r"^\s*(?:Q(?:uestion)?\s*)?(\d{1,3})\s*[.)번:]\s*(.+)$", re.I)
    answer_pattern = re.compile(r"^\s*(?:Answer|Ans|정답)\s*[:：]\s*(.+)$", re.I)
    choice_pattern = re.compile(r"^\s*(?:[A-Ea-e][.)]|[①②③④⑤])\s*(.+)$")
    blocks: list[dict] = []
    current: dict | None = None
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if match := question_pattern.match(line):
            if current and current.get("answer"):
                blocks.append(current)
            current = {"prompt": match.group(2).strip(), "choices": [], "answer": "", "explanation": ""}
            continue
        if current and (match := answer_pattern.match(line)):
            current["answer"] = match.group(1).strip()
            continue
        if current and (match := choice_pattern.match(line)):
            current["choices"].append(match.group(1).strip())
            continue
        if current:
            if current.get("answer"):
                current["explanation"] = f"{current['explanation']} {line}".strip()
            else:
                current["prompt"] = f"{current['prompt']} {line}".strip()
    if current and current.get("answer"):
        blocks.append(current)
    return [dict(item, position=index) for index, item in enumerate(blocks, start=1)]


def detect_question_count(text: str) -> int:
    parsed = parse_reading_quiz_sections(text) or parse_questions_from_text(text)
    if parsed:
        return len(parsed)
    numbers = [int(value) for value in re.findall(r"(?m)^\s*(\d{1,3})\s*[.)번]", text)]
    return max(numbers, default=0)


def insert_questions(db: sqlite3.Connection, content_id: int, questions: list[dict]) -> None:
    db.executemany(
        "INSERT INTO questions(content_id, position, prompt, choices_json, answer, explanation) VALUES(?,?,?,?,?,?)",
        [
            (
                content_id,
                item["position"],
                item["prompt"],
                json.dumps(item.get("choices", []), ensure_ascii=False),
                item["answer"],
                item.get("explanation", ""),
            )
            for item in questions
        ],
    )


def row_to_dict(row: sqlite3.Row) -> dict:
    item = dict(row)
    item.pop("extracted_text", None)
    item.pop("file_path", None)
    return item


def get_content(content_id: int) -> dict | None:
    with connect() as db:
        row = db.execute("SELECT * FROM contents WHERE id=?", (content_id,)).fetchone()
        if not row:
            return None
        item = row_to_dict(row)
        question_rows = db.execute("SELECT * FROM questions WHERE content_id=? ORDER BY position", (content_id,)).fetchall()
        item["questions"] = [
            {
                "id": question["id"],
                "position": question["position"],
                "prompt": question["prompt"],
                "choices": json.loads(question["choices_json"]),
                "answer": question["answer"],
                "explanation": question["explanation"],
            }
            for question in question_rows
        ]
        return item


QUIZ_SCHEMA = {
    "type": "object",
    "properties": {
        "questions": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "prompt": {"type": "string"},
                    "choices": {"type": "array", "items": {"type": "string"}},
                    "answer": {"type": "string"},
                    "explanation": {"type": "string"},
                },
                "required": ["prompt", "choices", "answer", "explanation"],
                "additionalProperties": False,
            },
        }
    },
    "required": ["questions"],
    "additionalProperties": False,
}


def call_openai(prompt: str) -> list[dict]:
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY가 서버에 설정되지 않았습니다.")
    body = {
        "model": os.getenv("OPENAI_MODEL", "gpt-5.4-mini"),
        "input": prompt,
        "text": {"format": {"type": "json_schema", "name": "english_quiz", "strict": True, "schema": QUIZ_SCHEMA}},
    }
    request = urllib.request.Request(
        "https://api.openai.com/v1/responses",
        data=json.dumps(body).encode("utf-8"),
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=180) as response:
            payload = json.load(response)
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"OpenAI API 오류: {detail[:500]}") from exc
    output_text = payload.get("output_text", "")
    if not output_text:
        for output in payload.get("output", []):
            for content in output.get("content", []):
                if content.get("type") == "output_text":
                    output_text += content.get("text", "")
    data = json.loads(output_text)
    return [question for index, item in enumerate(data["questions"], start=1) if (question := normalize_question(item, index))]


class ArchiveHandler(SimpleHTTPRequestHandler):
    server_version = "MKEnglishArchive/1.0"

    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(ROOT), **kwargs)

    def send_json(self, payload: object, status: int = 200) -> None:
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(data)))
        self.send_header("Cache-Control", "no-store")
        self.end_headers()
        self.wfile.write(data)

    def admin_password(self) -> str:
        return os.getenv("ADMIN_PASSWORD", "").strip()

    def admin_token(self) -> str:
        cookie = self.headers.get("Cookie", "")
        for part in cookie.split(";"):
            key, separator, value = part.strip().partition("=")
            if separator and key == "mk_admin_session":
                return value
        return ""

    def is_admin(self) -> bool:
        # Local development stays convenient until a test password is configured.
        if not self.admin_password():
            return True
        token = self.admin_token()
        expires_at = ADMIN_SESSIONS.get(token)
        if not expires_at:
            return False
        if expires_at <= datetime.now(timezone.utc):
            ADMIN_SESSIONS.pop(token, None)
            return False
        return True

    def require_admin(self) -> bool:
        if self.is_admin():
            return True
        self.send_json({"error": "관리자 로그인이 필요합니다."}, 401)
        return False

    def read_json(self) -> dict:
        length = int(self.headers.get("Content-Length", "0"))
        if length > 30 * 1024 * 1024:
            raise ValueError("파일은 최대 20MB까지 업로드할 수 있습니다.")
        return json.loads(self.rfile.read(length) or b"{}")

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path == "/admin.html" and not self.is_admin():
            self.send_response(302)
            self.send_header("Location", "/admin-login.html")
            self.end_headers()
            return
        if parsed.path == "/api/admin/session":
            self.send_json({"authenticated": self.is_admin(), "password_enabled": bool(self.admin_password())})
            return
        if parsed.path.startswith("/data/"):
            self.send_json({"error": "직접 접근할 수 없습니다."}, 404)
            return
        if parsed.path == "/api/health":
            self.send_json({"ok": True, "ai_enabled": bool(os.getenv("OPENAI_API_KEY")), "database": str(DB_PATH.name)})
            return
        if parsed.path == "/api/contents":
            query = parse_qs(parsed.query)
            clauses, values = [], []
            for field in ("kind", "school", "book", "topic", "status", "source_role"):
                if query.get(field):
                    clauses.append(f"{field}=?")
                    values.append(query[field][0])
            if not self.is_admin():
                clauses.append("status='published'")
            sql = "SELECT * FROM contents"
            if clauses:
                sql += " WHERE " + " AND ".join(clauses)
            sql += " ORDER BY updated_at DESC"
            with connect() as db:
                items = [row_to_dict(row) for row in db.execute(sql, values).fetchall()]
            self.send_json({"items": items})
            return
        if match := re.fullmatch(r"/api/contents/(\d+)", parsed.path):
            item = get_content(int(match.group(1)))
            if item and item["status"] != "published" and not self.is_admin():
                item = None
            self.send_json(item or {"error": "찾을 수 없습니다."}, 200 if item else 404)
            return
        super().do_GET()

    def do_POST(self) -> None:
        try:
            payload = self.read_json()
            if self.path == "/api/admin/login":
                expected = self.admin_password()
                supplied = str(payload.get("password") or "")
                if not expected or not hmac.compare_digest(supplied, expected):
                    self.send_json({"error": "비밀번호가 맞지 않습니다."}, 401)
                    return
                token = uuid.uuid4().hex + uuid.uuid4().hex
                ADMIN_SESSIONS[token] = datetime.now(timezone.utc) + timedelta(hours=ADMIN_SESSION_HOURS)
                data = json.dumps({"ok": True}, ensure_ascii=False).encode("utf-8")
                self.send_response(200)
                self.send_header("Content-Type", "application/json; charset=utf-8")
                self.send_header("Content-Length", str(len(data)))
                self.send_header("Cache-Control", "no-store")
                self.send_header("Set-Cookie", f"mk_admin_session={token}; Path=/; HttpOnly; SameSite=Lax; Max-Age={ADMIN_SESSION_HOURS * 3600}")
                self.end_headers()
                self.wfile.write(data)
                return
            if self.path == "/api/admin/logout":
                ADMIN_SESSIONS.pop(self.admin_token(), None)
                data = b'{"ok":true}'
                self.send_response(200)
                self.send_header("Content-Type", "application/json; charset=utf-8")
                self.send_header("Content-Length", str(len(data)))
                self.send_header("Set-Cookie", "mk_admin_session=; Path=/; HttpOnly; SameSite=Lax; Max-Age=0")
                self.end_headers()
                self.wfile.write(data)
                return
            if not self.require_admin():
                return
            if self.path == "/api/upload":
                self.handle_upload(payload)
                return
            if self.path == "/api/generate":
                self.handle_generate(payload)
                return
            if match := re.fullmatch(r"/api/contents/(\d+)/publish", self.path):
                with connect() as db:
                    db.execute("UPDATE contents SET status='published', updated_at=? WHERE id=?", (now(), int(match.group(1))))
                self.send_json(get_content(int(match.group(1))))
                return
            self.send_json({"error": "지원하지 않는 API입니다."}, 404)
        except (ValueError, json.JSONDecodeError) as exc:
            self.send_json({"error": str(exc)}, 400)
        except RuntimeError as exc:
            self.send_json({"error": str(exc)}, 503)
        except Exception as exc:
            self.send_json({"error": f"서버 오류: {exc}"}, 500)

    def do_PUT(self) -> None:
        try:
            if not self.require_admin():
                return
            match = re.fullmatch(r"/api/contents/(\d+)", self.path)
            if not match:
                self.send_json({"error": "지원하지 않는 API입니다."}, 404)
                return
            payload = self.read_json()
            allowed = {key: payload[key] for key in ("title", "status", "school", "topic", "book") if key in payload}
            if not allowed:
                raise ValueError("수정할 값이 없습니다.")
            assignments = ", ".join(f"{key}=?" for key in allowed)
            with connect() as db:
                db.execute(f"UPDATE contents SET {assignments}, updated_at=? WHERE id=?", (*allowed.values(), now(), int(match.group(1))))
            self.send_json(get_content(int(match.group(1))))
        except ValueError as exc:
            self.send_json({"error": str(exc)}, 400)

    def handle_upload(self, payload: dict) -> None:
        filename = safe_name(str(payload.get("filename", "upload.bin")))
        raw = base64.b64decode(payload.get("data_base64", ""), validate=True)
        if not raw:
            raise ValueError("빈 파일입니다.")
        if len(raw) > 20 * 1024 * 1024:
            raise ValueError("파일은 최대 20MB까지 업로드할 수 있습니다.")
        stored_path = UPLOAD_DIR / f"{uuid.uuid4().hex}_{filename}"
        stored_path.write_bytes(raw)
        mime_type = str(payload.get("mime_type") or mimetypes.guess_type(filename)[0] or "application/octet-stream")
        text = extract_text(stored_path, mime_type)
        kind = str(payload.get("kind") or "source")
        title = str(payload.get("title") or Path(filename).stem).strip()
        school = str(payload.get("school") or "").strip()
        topic = str(payload.get("topic") or "").strip()
        book = str(payload.get("book") or "").strip()
        source_role = str(payload.get("source_role") or "").strip()
        questions = parse_structured_questions(text, stored_path.suffix.lower()) if text else []
        detected_count = len(questions) or detect_question_count(text)
        status = "published" if kind == "novel_quiz" and questions else "draft"
        with connect() as db:
            previous = db.execute(
                "SELECT id, version FROM contents WHERE kind=? AND title=? AND school=? AND topic=? AND book=? ORDER BY version DESC LIMIT 1",
                (kind, title, school, topic, book),
            ).fetchone()
            version = (previous["version"] + 1) if previous else 1
            cursor = db.execute(
                """INSERT INTO contents(kind,title,school,topic,book,source_role,filename,mime_type,file_path,extracted_text,status,question_count,version,parent_id,created_at,updated_at)
                   VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (kind, title, school, topic, book, source_role, filename, mime_type, str(stored_path), text, status, detected_count, version, previous["id"] if previous else None, now(), now()),
            )
            content_id = cursor.lastrowid
            insert_questions(db, content_id, questions)
        self.send_json({"item": get_content(content_id), "extracted_characters": len(text), "parsed_questions": len(questions)}, 201)

    def handle_generate(self, payload: dict) -> None:
        generation_type = str(payload.get("type") or "")
        school = str(payload.get("school") or "")
        topic = str(payload.get("topic") or "")
        with connect() as db:
            if generation_type == "mock_exam":
                exam = db.execute("SELECT * FROM contents WHERE school=? AND source_role='past_exam' ORDER BY updated_at DESC LIMIT 1", (school,)).fetchone()
                textbook = db.execute("SELECT * FROM contents WHERE school=? AND source_role='textbook_range' ORDER BY updated_at DESC LIMIT 1", (school,)).fetchone()
                if not exam or not textbook:
                    raise ValueError("기출 시험지와 교과서 시험범위 자료를 모두 업로드해 주세요.")
                count = exam["question_count"] or detect_question_count(exam["extracted_text"])
                if not count:
                    raise ValueError("기출 PDF에서 문항 수를 인식하지 못했습니다.")
                prompt = f"""한국 중고등학교 영어 내신 모의고사를 정확히 {count}문항 작성하세요.
기출 시험지는 문항 유형, 난이도, 배점 구조만 참고하고 문구를 복제하지 마세요.
새 문제의 내용은 반드시 아래 교과서 시험범위 지문에서 출제하세요.
각 문항은 prompt, choices, answer, explanation을 포함해야 합니다.

[기출 시험 구조]
{exam['extracted_text'][:60000]}

[이번 시험범위 교과서 지문]
{textbook['extracted_text'][:60000]}"""
                title = f"{school} 기출유형 모의고사"
                kind, book = "mock_exam", ""
            elif generation_type == "grammar":
                source = db.execute("SELECT * FROM contents WHERE topic=? ORDER BY updated_at DESC LIMIT 1", (topic,)).fetchone()
                if not source:
                    raise ValueError("해당 문법 아이템의 소스 자료를 먼저 업로드해 주세요.")
                count = 15
                prompt = f"""중학생 영어 학습용 {topic} 문법 퀴즈를 정확히 15문항 작성하세요.
객관식과 빈칸 문제를 섞고, prompt, choices, answer, explanation을 포함하세요.
아래 프린트의 학습 범위와 난이도를 따르되 문제 문구는 새로 작성하세요.

{source['extracted_text'][:80000]}"""
                title = f"{topic} Grammar Quiz"
                kind, school, book = "grammar_quiz", "", ""
            else:
                raise ValueError("지원하지 않는 생성 유형입니다.")
            questions = call_openai(prompt)
            if len(questions) != count:
                raise RuntimeError(f"생성 문항 수가 예상과 다릅니다: {len(questions)}/{count}")
            cursor = db.execute(
                """INSERT INTO contents(kind,title,school,topic,book,status,question_count,version,created_at,updated_at)
                   VALUES(?,?,?,?,?,'review',?,1,?,?)""",
                (kind, title, school, topic, book, count, now(), now()),
            )
            content_id = cursor.lastrowid
            insert_questions(db, content_id, questions)
        self.send_json({"item": get_content(content_id)}, 201)


def run() -> None:
    init_db()
    port = int(os.getenv("PORT", "4173"))
    server = ThreadingHTTPServer(("127.0.0.1", port), ArchiveHandler)
    print(f"MK ENGLISH ARCHIVE running at http://127.0.0.1:{port}")
    print(f"AI generation: {'enabled' if os.getenv('OPENAI_API_KEY') else 'disabled (set OPENAI_API_KEY)'}")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        pass
    finally:
        server.server_close()


if __name__ == "__main__":
    run()
